"""
Document Processors — extraction de fichiers texte structurés.

Formats : PDF, DOCX, TXT, Markdown
Chaque processeur retourne : list[{type: "text"/"table", content: str, page: int}]
"""
import os
import re
import logging

logger = logging.getLogger(__name__)


# ══════════════════════════════════════════════════════
# PDF — pdfplumber + PyPDF fallback
# ══════════════════════════════════════════════════════

def extract_pdf(file_path: str) -> list[dict]:
    """
    Extrait texte + tableaux d'un PDF.
    Utilisation : politique_rh.pdf, contrat.pdf, rapport_financier.pdf
    3 méthodes de détection de tableaux.
    """
    import pdfplumber

    content_blocks = []

    with pdfplumber.open(file_path) as pdf:
        for page_num, page in enumerate(pdf.pages, 1):
            # ── Tableaux ──
            tables_found = []
            for strategy in [
                {"vertical_strategy": "lines", "horizontal_strategy": "lines"},
                {"vertical_strategy": "text", "horizontal_strategy": "text"},
                {"vertical_strategy": "text", "horizontal_strategy": "lines"},
            ]:
                try:
                    tables = page.extract_tables(table_settings=strategy)
                    if tables:
                        tables_found = tables
                        break
                except Exception:
                    continue

            for table in tables_found:
                md = _table_to_markdown(table)
                if md:
                    content_blocks.append({"type": "table", "content": md, "page": page_num})

            # ── Texte ──
            try:
                text = page.extract_text()
                if text and text.strip() and len(text.strip()) > 20:
                    content_blocks.append({"type": "text", "content": text.strip(), "page": page_num})
            except Exception:
                continue

    # Fallback PyPDF
    if not content_blocks:
        try:
            from pypdf import PdfReader
            reader = PdfReader(file_path)
            for i, page in enumerate(reader.pages, 1):
                text = page.extract_text()
                if text and text.strip():
                    content_blocks.append({"type": "text", "content": text.strip(), "page": i})
        except Exception as e:
            logger.error(f"PDF fallback failed: {e}")

    return content_blocks


# ══════════════════════════════════════════════════════
# DOCX — python-docx
# ══════════════════════════════════════════════════════

def extract_docx(file_path: str) -> list[dict]:
    """
    Extrait paragraphes + tableaux d'un Word.
    Utilisation : procedure_teletravail.docx, guide_installation.docx
    Les Headings sont détectés et préfixés [Section: titre].
    """
    from docx import Document

    doc = Document(file_path)
    content_blocks = []
    current_section = ""
    current_lines = []
    page_estimate = 1

    for element in doc.element.body:
        tag = element.tag.split("}")[-1]

        if tag == "p":
            para = None
            for p in doc.paragraphs:
                if p._element is element:
                    para = p
                    break
            if para is None:
                continue

            text = para.text.strip()
            if not text:
                continue

            style_name = para.style.name if para.style else ""

            if "Heading" in style_name or "heading" in style_name:
                # Sauvegarder le bloc précédent
                if current_lines:
                    block = "\n".join(current_lines)
                    if current_section:
                        block = f"[Section: {current_section}]\n{block}"
                    content_blocks.append({"type": "text", "content": block, "page": page_estimate})
                    current_lines = []
                current_section = text
                total_chars = sum(len(b["content"]) for b in content_blocks)
                page_estimate = max(1, total_chars // 3000 + 1)
            else:
                current_lines.append(text)

        elif tag == "tbl":
            # Sauvegarder avant le tableau
            if current_lines:
                block = "\n".join(current_lines)
                if current_section:
                    block = f"[Section: {current_section}]\n{block}"
                content_blocks.append({"type": "text", "content": block, "page": page_estimate})
                current_lines = []

            for table in doc.tables:
                if table._element is element:
                    rows = []
                    for row in table.rows:
                        cells = [cell.text.strip() for cell in row.cells]
                        rows.append(cells)
                    md = _rows_to_markdown(rows)
                    if md:
                        content_blocks.append({"type": "table", "content": md, "page": page_estimate})
                    break

    # Dernier bloc
    if current_lines:
        block = "\n".join(current_lines)
        if current_section:
            block = f"[Section: {current_section}]\n{block}"
        content_blocks.append({"type": "text", "content": block, "page": page_estimate})

    return content_blocks


# ══════════════════════════════════════════════════════
# TXT — texte brut avec détection d'encodage
# ══════════════════════════════════════════════════════

def extract_txt(file_path: str) -> list[dict]:
    """
    Lit un fichier texte brut.
    Utilisation : faq_employes.txt, notes_reunion.txt, changelog.txt
    Détecte l'encodage automatiquement (UTF-8, Latin-1, Windows-1256).
    """
    import chardet

    with open(file_path, "rb") as f:
        raw = f.read()
        detected = chardet.detect(raw)
        encoding = detected.get("encoding", "utf-8") or "utf-8"

    text = raw.decode(encoding, errors="replace").strip()
    if not text:
        return []

    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
    return [{"type": "text", "content": para, "page": 1} for para in paragraphs]


# ══════════════════════════════════════════════════════
# MARKDOWN — fichiers .md
# ══════════════════════════════════════════════════════

def extract_markdown(file_path: str) -> list[dict]:
    """
    Lit un fichier Markdown, découpe par sections (#).
    Utilisation : README.md, api_documentation.md, guide_onboarding.md
    """
    import chardet

    with open(file_path, "rb") as f:
        raw = f.read()
        detected = chardet.detect(raw)
        encoding = detected.get("encoding", "utf-8") or "utf-8"

    text = raw.decode(encoding, errors="replace").strip()
    if not text:
        return []

    content_blocks = []
    current_section = ""
    current_lines = []

    for line in text.split("\n"):
        stripped = line.strip()
        if stripped.startswith("#"):
            if current_lines:
                block = "\n".join(current_lines).strip()
                if block:
                    if current_section:
                        block = f"[Section: {current_section}]\n{block}"
                    content_blocks.append({"type": "text", "content": block, "page": 1})
                current_lines = []
            current_section = stripped.lstrip("#").strip()
        else:
            if stripped:
                current_lines.append(stripped)

    if current_lines:
        block = "\n".join(current_lines).strip()
        if block:
            if current_section:
                block = f"[Section: {current_section}]\n{block}"
            content_blocks.append({"type": "text", "content": block, "page": 1})

    if not content_blocks and text:
        paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
        for para in paragraphs:
            content_blocks.append({"type": "text", "content": para, "page": 1})

    return content_blocks


# ══════════════════════════════════════════════════════
# HELPERS
# ══════════════════════════════════════════════════════

def _table_to_markdown(table) -> str | None:
    """Convertit un tableau pdfplumber en Markdown."""
    if not table or len(table) < 2:
        return None
    try:
        headers = [str(h).strip() if h else "" for h in table[0]]
        md_lines = ["| " + " | ".join(headers) + " |"]
        md_lines.append("| " + " | ".join(["---"] * len(headers)) + " |")
        for row in table[1:]:
            cells = [str(c).strip() if c else "" for c in row]
            if len(cells) == len(headers):
                md_lines.append("| " + " | ".join(cells) + " |")
        result = "\n".join(md_lines)
        return result if len(result) > 20 else None
    except Exception:
        return None


def _rows_to_markdown(rows: list[list[str]]) -> str | None:
    """Convertit une liste de rows en Markdown."""
    if len(rows) < 2:
        return None
    headers = rows[0]
    md_lines = ["| " + " | ".join(headers) + " |"]
    md_lines.append("| " + " | ".join(["---"] * len(headers)) + " |")
    for row in rows[1:]:
        if len(row) == len(headers):
            md_lines.append("| " + " | ".join(row) + " |")
    result = "\n".join(md_lines)
    return result if len(result) > 20 else None
